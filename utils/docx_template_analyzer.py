#!/usr/bin/env python3
"""
DOCX Template Analyzer for Section E Templates

This script analyzes .docx templates to understand their structure, 
placeholders, and formatting elements.
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def analyze_docx_template(template_path: str):
    """
    Analyze DOCX template structure and extract placeholder information
    
    Args:
        template_path (str): Path to the DOCX template file
    """
    if not os.path.exists(template_path):
        print(f"❌ Template not found: {template_path}")
        return None
    
    try:
        doc = Document(template_path)
        
        print(f"📄 Analyzing DOCX template: {template_path}")
        print("=" * 60)
        
        # Basic document info
        print(f"📊 Document Structure:")
        print(f"  - Paragraphs: {len(doc.paragraphs)}")
        print(f"  - Tables: {len(doc.tables)}")
        print(f"  - Sections: {len(doc.sections)}")
        
        # Page setup info
        section = doc.sections[0] if doc.sections else None
        print(f"\n📐 Page Setup:")
        if section and hasattr(section, 'page_width') and section.page_width:
            print(f"  - Width: {section.page_width.inches:.2f} inches")
            print(f"  - Height: {section.page_height.inches:.2f} inches")
            print(f"  - Top margin: {section.top_margin.inches:.2f} inches")
            print(f"  - Bottom margin: {section.bottom_margin.inches:.2f} inches")
            print(f"  - Left margin: {section.left_margin.inches:.2f} inches")
            print(f"  - Right margin: {section.right_margin.inches:.2f} inches")
        else:
            print("  - No sections or page setup found")
        
        # Analyze paragraphs for placeholders
        placeholders = set()
        print(f"\n📝 Paragraph Analysis:")
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                print(f"  {i+1:2d}: {text[:80]}{'...' if len(text) > 80 else ''}")
                
                # Look for placeholder patterns
                import re
                # Common patterns: {{name}}, [name], {name}, <<name>>, etc.
                patterns = [
                    r'\{\{([^}]+)\}\}',  # {{placeholder}}
                    r'\[([^\]]+)\]',     # [placeholder]
                    r'\{([^}]+)\}',      # {placeholder}
                    r'<<([^>]+)>>',      # <<placeholder>>
                    r'__([^_]+)__',      # __placeholder__
                ]
                
                for pattern in patterns:
                    matches = re.findall(pattern, text)
                    for match in matches:
                        placeholders.add(match.strip())
        
        # Analyze tables
        print(f"\n📋 Table Analysis:")
        for table_idx, table in enumerate(doc.tables):
            print(f"  Table {table_idx + 1}: {len(table.rows)} rows × {len(table.columns)} columns")
            
            # Check table cells for placeholders
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        # Look for placeholders in cell text
                        import re
                        patterns = [
                            r'\{\{([^}]+)\}\}',
                            r'\[([^\]]+)\]',
                            r'\{([^}]+)\}',
                            r'<<([^>]+)>>',
                            r'__([^_]+)__',
                        ]
                        
                        for pattern in patterns:
                            matches = re.findall(pattern, cell_text)
                            for match in matches:
                                placeholders.add(match.strip())
                        
                        if row_idx < 3 or any(p in cell_text.lower() for p in ['name', 'role', 'experience', 'education']):
                            print(f"    Row {row_idx+1}, Col {col_idx+1}: {cell_text[:60]}{'...' if len(cell_text) > 60 else ''}")
        
        # Print found placeholders
        print(f"\n🏷️  Found Placeholders:")
        if placeholders:
            for placeholder in sorted(placeholders):
                print(f"  - {placeholder}")
        else:
            print("  No standard placeholders found")
            print("  💡 This template may use a different placeholder format or be ready-made content")
        
        # Analyze styles
        print(f"\n🎨 Style Information:")
        try:
            styles = doc.styles
            print(f"  - Available styles: {len(list(styles))}")
            
            # Show some key styles
            key_styles = ['Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'Title']
            for style_name in key_styles:
                try:
                    style = styles[style_name]
                    print(f"    - {style_name}: Available")
                except:
                    print(f"    - {style_name}: Not found")
        except Exception as e:
            print(f"  Error analyzing styles: {e}")
        
        page_width = 0
        page_height = 0
        if section and hasattr(section, 'page_width') and section.page_width:
            page_width = section.page_width.inches
            page_height = section.page_height.inches
        
        return {
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables),
            'placeholders': list(placeholders),
            'page_width': page_width,
            'page_height': page_height
        }
        
    except Exception as e:
        print(f"❌ Error analyzing template: {e}")
        return None


def suggest_placeholder_mapping():
    """Suggest common placeholder mappings for Section E templates"""
    print(f"\n💡 Suggested Placeholder Mappings for Section E:")
    print("=" * 60)
    
    mappings = {
        "Employee Information": {
            "{{employee_name}}": "employee_data['name']",
            "{{role_in_contract}}": "employee_data['role_in_contract']", 
            "{{total_experience}}": "employee_data['years_experience']['total']",
            "{{current_firm_experience}}": "employee_data['years_experience']['with_current_firm']",
            "{{firm_name_location}}": "employee_data['firm_name_and_location']",
            "{{education}}": "employee_data['education']",
            "{{professional_registration}}": "employee_data['current_professional_registration']",
            "{{other_qualifications}}": "employee_data['other_professional_qualifications']"
        },
        "Project Information": {
            "{{project_title_1}}": "projects[0]['title_and_location']",
            "{{project_description_1}}": "projects[0]['description']['scope']",
            "{{project_year_1}}": "projects[0]['year_completed']['professional_services']",
            "{{project_cost_1}}": "projects[0]['description']['cost']",
            "{{project_role_1}}": "projects[0]['description']['role']"
        }
    }
    
    for section, mapping in mappings.items():
        print(f"\n{section}:")
        for placeholder, data_path in mapping.items():
            print(f"  {placeholder:<30} → {data_path}")


def main():
    """Main function"""
    print("🚀 DOCX Template Analyzer for Section E")
    print("=" * 60)
    
    # Check command line arguments
    if len(sys.argv) > 1:
        template_path = sys.argv[1]
    else:
        template_path = "templates/Template Section E Resume.docx"
    
    # Analyze the template
    result = analyze_docx_template(template_path)
    
    if result:
        # Suggest placeholder mappings
        suggest_placeholder_mapping()
        
        print(f"\n✅ Analysis complete!")
        print(f"📋 Summary: {result['paragraphs']} paragraphs, {result['tables']} tables, {len(result['placeholders'])} placeholders")
    else:
        print(f"\n❌ Analysis failed")


if __name__ == "__main__":
    main() 