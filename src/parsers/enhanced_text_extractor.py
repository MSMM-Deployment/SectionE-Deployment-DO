#!/usr/bin/env python3
"""
Enhanced PDF Text Extractor with Encoding Fixes

This module provides multiple extraction methods and text preprocessing
to handle various PDF encoding issues before sending to OpenAI.
"""

import re
import unicodedata
import logging
from typing import List, Dict, Optional, Tuple, Any
from pathlib import Path

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from PyPDF2 import PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    from pdfminer.high_level import extract_text
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False

try:
    import subprocess
    import shutil
    SUBPROCESS_AVAILABLE = True
    # Check if antiword is available (for legacy .doc files)
    ANTIWORD_AVAILABLE = shutil.which('antiword') is not None
    # Check if catdoc is available (alternative for .doc files)
    CATDOC_AVAILABLE = shutil.which('catdoc') is not None
except ImportError:
    SUBPROCESS_AVAILABLE = False
    ANTIWORD_AVAILABLE = False
    CATDOC_AVAILABLE = False

try:
    import textract
    TEXTRACT_AVAILABLE = True
except ImportError:
    TEXTRACT_AVAILABLE = False


class EnhancedTextExtractor:
    """Enhanced text extractor for PDFs and Word documents with encoding fixes"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
        # CID to Unicode mapping for common characters
        self.cid_mappings = {
            # Common CID patterns found in government forms
            "(cid:20)": "1",   # 1
            "(cid:21)": "2",   # 2  
            "(cid:22)": "3",   # 3
            "(cid:23)": "4",   # 4
            "(cid:24)": "5",   # 5
            "(cid:25)": "6",   # 6
            "(cid:26)": "7",   # 7
            "(cid:27)": "8",   # 8
            "(cid:28)": "9",   # 9
            "(cid:19)": "0",   # 0
            "(cid:17)": ".",   # period
            "(cid:3)": " ",    # space
            "(cid:36)": "A",   # A
            "(cid:37)": "B",   # B
            "(cid:38)": "C",   # C
            "(cid:39)": "D",   # D
            "(cid:40)": "E",   # E
            "(cid:41)": "F",   # F
            "(cid:42)": "G",   # G
            "(cid:43)": "H",   # H
            "(cid:44)": "I",   # I
            "(cid:45)": "J",   # J
            "(cid:46)": "K",   # K
            "(cid:47)": "L",   # L
            "(cid:48)": "M",   # M
            "(cid:49)": "N",   # N
            "(cid:50)": "O",   # O
            "(cid:51)": "P",   # P
            "(cid:52)": "Q",   # Q
            "(cid:53)": "R",   # R
            "(cid:54)": "S",   # S
            "(cid:55)": "T",   # T
            "(cid:56)": "U",   # U
            "(cid:57)": "V",   # V
            "(cid:58)": "W",   # W
            "(cid:59)": "X",   # X
            "(cid:60)": "Y",   # Y
            "(cid:61)": "Z",   # Z
            "(cid:68)": "a",   # a
            "(cid:69)": "b",   # b
            "(cid:70)": "c",   # c
            "(cid:71)": "d",   # d
            "(cid:72)": "e",   # e
            "(cid:73)": "f",   # f
            "(cid:74)": "g",   # g
            "(cid:75)": "h",   # h
            "(cid:76)": "i",   # i
            "(cid:77)": "j",   # j
            "(cid:78)": "k",   # k
            "(cid:79)": "l",   # l
            "(cid:80)": "m",   # m
            "(cid:81)": "n",   # n
            "(cid:82)": "o",   # o
            "(cid:83)": "p",   # p
            "(cid:84)": "q",   # q
            "(cid:85)": "r",   # r
            "(cid:86)": "s",   # s
            "(cid:87)": "t",   # t
            "(cid:88)": "u",   # u
            "(cid:89)": "v",   # v
            "(cid:90)": "w",   # w
            "(cid:91)": "x",   # x
            "(cid:92)": "y",   # y
            "(cid:93)": "z",   # z
        }
        
        # Character replacement mappings for garbled text
        self.char_replacements = {
            # Common garbled patterns from PDF extraction
            "1$0(": "NAME",
            "52/(": "ROLE", 
            ",17+,6": "IN THIS",
            "&2175$&7": "CONTRACT",
            "<($56": "YEARS",
            "(;3(5,(1&(": "EXPERIENCE",
            "727$/": "TOTAL",
            ":,7+": "WITH",
            "&855(17": "CURRENT",
            "),50": "FIRM",
            "),501$0(": "FIRM NAME",
            "$1'": "AND",
            "/2&$7,21": "LOCATION",
            "('8&$7,21": "EDUCATION",
            "&855(17": "CURRENT",
            "352)(66,21$/": "PROFESSIONAL",
            "5(*,675$7,21": "REGISTRATION",
            "%6": "BS",
            "&LYLO": "Civil",
            "(QJLQHHULQJ": "Engineering",
            "0LFKLJDQ": "Michigan",
            "7HFKQRORJLFDO": "Technological",
        }

    def extract_text_pymupdf(self, pdf_path: str) -> str:
        """Extract text using PyMuPDF (best for most PDFs)"""
        if not PYMUPDF_AVAILABLE:
            raise ImportError("PyMuPDF not available")
        
        try:
            doc = fitz.open(pdf_path)
            text = ""
            for page in doc:
                # PyMuPDF page text extraction
                page_text = getattr(page, 'get_text', lambda *args: '')()
                text += page_text + "\n"
            doc.close()
            return text.strip()
        except Exception as e:
            self.logger.error(f"PyMuPDF extraction failed: {e}")
            return ""

    def extract_text_pdfplumber(self, pdf_path: str) -> str:
        """Extract text using pdfplumber"""
        if not PDFPLUMBER_AVAILABLE:
            raise ImportError("pdfplumber not available")
        
        try:
            text = ""
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text.strip()
        except Exception as e:
            self.logger.error(f"pdfplumber extraction failed: {e}")
            return ""

    def extract_text_pypdf2(self, pdf_path: str) -> str:
        """Extract text using PyPDF2"""
        if not PYPDF2_AVAILABLE:
            raise ImportError("PyPDF2 not available")
        
        try:
            text = ""
            with open(pdf_path, 'rb') as file:
                reader = PdfReader(file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text.strip()
        except Exception as e:
            self.logger.error(f"PyPDF2 extraction failed: {e}")
            return ""

    def extract_text_pdfminer(self, pdf_path: str) -> str:
        """Extract text using pdfminer"""
        if not PDFMINER_AVAILABLE:
            raise ImportError("pdfminer not available")
        
        try:
            return extract_text(pdf_path).strip()
        except Exception as e:
            self.logger.error(f"pdfminer extraction failed: {e}")
            return ""

    def extract_text_python_docx(self, doc_path: str) -> str:
        """Extract text using python-docx (for .docx files)"""
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx not available")
        
        try:
            doc = Document(doc_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return "\n".join(text_parts).strip()
            
        except Exception as e:
            self.logger.error(f"python-docx extraction failed: {e}")
            return ""

    def extract_text_docx2txt(self, doc_path: str) -> str:
        """Extract text using docx2txt (simpler, for .docx files)"""
        if not DOCX2TXT_AVAILABLE:
            raise ImportError("docx2txt not available")
        
        try:
            text = docx2txt.process(doc_path)
            return text.strip() if text else ""
        except Exception as e:
            self.logger.error(f"docx2txt extraction failed: {e}")
            return ""

    def extract_text_antiword(self, doc_path: str) -> str:
        """Extract text using antiword (for legacy .doc files)"""
        if not ANTIWORD_AVAILABLE:
            raise ImportError("antiword not available")
        
        try:
            result = subprocess.run(
                ['antiword', doc_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0:
                return result.stdout.strip()
            else:
                self.logger.error(f"antiword failed with return code {result.returncode}: {result.stderr}")
                return ""
        except subprocess.TimeoutExpired:
            self.logger.error("antiword extraction timed out")
            return ""
        except Exception as e:
            self.logger.error(f"antiword extraction failed: {e}")
            return ""

    def extract_text_catdoc(self, doc_path: str) -> str:
        """Extract text using catdoc (alternative for legacy .doc files)"""
        if not CATDOC_AVAILABLE:
            raise ImportError("catdoc not available")
        
        try:
            result = subprocess.run(
                ['catdoc', doc_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0:
                return result.stdout.strip()
            else:
                self.logger.error(f"catdoc failed with return code {result.returncode}: {result.stderr}")
                return ""
        except subprocess.TimeoutExpired:
            self.logger.error("catdoc extraction timed out")
            return ""
        except Exception as e:
            self.logger.error(f"catdoc extraction failed: {e}")
            return ""

    def extract_text_textract(self, doc_path: str) -> str:
        """Extract text using textract (comprehensive document processing)"""
        if not TEXTRACT_AVAILABLE:
            raise ImportError("textract not available")
        
        try:
            text = textract.process(doc_path).decode('utf-8')
            return text.strip() if text else ""
        except Exception as e:
            self.logger.error(f"textract extraction failed: {e}")
            return ""

    def decode_cid_patterns(self, text: str) -> str:
        """Decode CID patterns to readable text"""
        decoded_text = text
        
        # Replace known CID mappings
        for cid, char in self.cid_mappings.items():
            decoded_text = decoded_text.replace(cid, char)
        
        # Remove any remaining unknown CID patterns
        decoded_text = re.sub(r'\(cid:\d+\)', '', decoded_text)
        
        return decoded_text

    def fix_garbled_text(self, text: str) -> str:
        """Fix common garbled text patterns"""
        fixed_text = text
        
        # Apply character replacements
        for garbled, fixed in self.char_replacements.items():
            fixed_text = fixed_text.replace(garbled, fixed)
        
        return fixed_text

    def normalize_unicode(self, text: str) -> str:
        """Normalize unicode characters"""
        try:
            # Normalize unicode to canonical form
            normalized = unicodedata.normalize('NFKC', text)
            
            # Remove or replace problematic characters
            normalized = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', normalized)
            
            return normalized
        except Exception as e:
            self.logger.warning(f"Unicode normalization failed: {e}")
            return text

    def clean_extracted_text(self, text: str) -> str:
        """Apply all text cleaning and preprocessing"""
        if not text:
            return ""
        
        # Step 1: Decode CID patterns
        text = self.decode_cid_patterns(text)
        
        # Step 2: Fix garbled text patterns  
        text = self.fix_garbled_text(text)
        
        # Step 3: Normalize unicode
        text = self.normalize_unicode(text)
        
        # Step 4: Clean up whitespace
        text = re.sub(r'\s+', ' ', text)  # Multiple spaces to single
        text = re.sub(r'\n\s*\n', '\n\n', text)  # Multiple newlines to double
        
        # Step 5: Remove excessive blank lines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()

    def extract_text_with_fallback(self, file_path: str) -> Tuple[str, str]:
        """
        Extract text using multiple methods with fallback for PDFs and Word documents
        
        Returns:
            Tuple[str, str]: (extracted_text, method_used)
        """
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            methods = [
                ("PyMuPDF", self.extract_text_pymupdf),
                ("PyPDF2", self.extract_text_pypdf2),
                ("pdfplumber", self.extract_text_pdfplumber), 
                ("pdfminer", self.extract_text_pdfminer),
            ]
        elif file_ext == '.docx':
            # Modern Word format (.docx) - use Office Open XML methods
            methods = [
                ("python-docx", self.extract_text_python_docx),
                ("docx2txt", self.extract_text_docx2txt),
                ("textract", self.extract_text_textract),
            ]
        elif file_ext == '.doc':
            # Legacy Word format (.doc) - use binary format methods
            methods = [
                ("antiword", self.extract_text_antiword),
                ("catdoc", self.extract_text_catdoc),
                ("textract", self.extract_text_textract),
                ("python-docx", self.extract_text_python_docx),  # Sometimes works
                ("docx2txt", self.extract_text_docx2txt),        # Sometimes works
            ]
        else:
            self.logger.error(f"Unsupported file format: {file_ext}")
            return "", "unsupported"
        
        for method_name, method_func in methods:
            try:
                text = method_func(file_path)
                if text and len(text.strip()) > 50:  # Minimum viable text (lower for Word docs)
                    cleaned_text = self.clean_extracted_text(text)
                    self.logger.info(f"Successfully extracted text using {method_name}")
                    return cleaned_text, method_name
            except ImportError:
                self.logger.warning(f"{method_name} not available")
                continue
            except Exception as e:
                self.logger.warning(f"{method_name} failed: {e}")
                continue
        
        return "", "none"

    def analyze_text_quality(self, text: str) -> Dict[str, Any]:
        """Analyze the quality of extracted text"""
        if not text:
            return {"quality": "empty", "issues": ["No text extracted"]}
        
        issues = []
        
        # Check for CID patterns
        cid_count = len(re.findall(r'\(cid:\d+\)', text))
        if cid_count > 0:
            issues.append(f"Contains {cid_count} CID patterns")
        
        # Check for garbled text patterns
        garbled_patterns = sum(1 for pattern in self.char_replacements.keys() if pattern in text)
        if garbled_patterns > 0:
            issues.append(f"Contains {garbled_patterns} garbled patterns")
        
        # Check for character encoding issues
        non_ascii = sum(1 for char in text if ord(char) > 127)
        if non_ascii > len(text) * 0.1:  # More than 10% non-ASCII
            issues.append(f"High non-ASCII character ratio: {non_ascii}/{len(text)}")
        
        # Determine quality
        if not issues:
            quality = "good"
        elif len(issues) <= 2:
            quality = "fair"
        else:
            quality = "poor"
        
        return {
            "quality": quality,
            "issues": issues,
            "length": len(text),
            "cid_patterns": cid_count,
            "garbled_patterns": garbled_patterns,
            "non_ascii_ratio": non_ascii / len(text) if text else 0
        }


def test_extraction(file_path: str) -> None:
    """Test the enhanced extraction on a specific file (PDF or Word)"""
    extractor = EnhancedTextExtractor()
    
    file_type = Path(file_path).suffix.lower()
    print(f"Testing extraction for: {file_path} ({file_type})")
    print("=" * 60)
    
    # Extract text
    text, method = extractor.extract_text_with_fallback(file_path)
    
    # Analyze quality
    quality = extractor.analyze_text_quality(text)
    
    print(f"Extraction method: {method}")
    print(f"Text quality: {quality['quality']}")
    print(f"Text length: {quality['length']}")
    print(f"Issues: {quality['issues']}")
    
    print("\nFirst 500 characters:")
    print("-" * 30)
    print(text[:500])
    
    if len(text) > 500:
        print("\n... (truncated)")


if __name__ == "__main__":
    # Test with the problematic PDF
    test_files = [
        "Sample-SectionE/Jim-Wilson-2.pdf",
        # Add test Word documents if available
    ]
    
    for test_file in test_files:
        if Path(test_file).exists():
            test_extraction(test_file)
            print("\n" + "="*60 + "\n")
        else:
            print(f"Test file not found: {test_file}") 