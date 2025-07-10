#!/usr/bin/env python3
"""
DOCX Section E Generator

This script generates Section E resumes using Word (.docx) templates,
providing much better formatting control than PDF templates.
Supports intelligent cell identification, text wrapping, and formatting.
"""

import os
import sys
import json
import re
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
# from docx.oxml.shared import OxmlElement, qn  # Not needed for current implementation

# Add project root to Python path for utils imports
project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.path.insert(0, project_root)

from utils.supabase_template_downloader import SupabaseTemplateDownloader


class DOCXSectionEGenerator:
    """Generate Section E resumes using DOCX templates"""
    
    def __init__(self, template_path: str = "templates/Template Section E Resume.docx"):
        """
        Initialize the DOCX generator
        
        Args:
            template_path (str): Path to the DOCX template file (local path or filename)
        """
        self.template_path = template_path
        self.output_dir = Path("OutputFiles/DOCX")
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize template downloader for cloud templates
        self.template_downloader = SupabaseTemplateDownloader()
        
        # Precise cell mapping based on template structure (31 rows × 7 columns)
        # These are the exact empty cells where data should go
        self.employee_cell_mapping = {
            'name': (2, 0),              # Row 3, Col 1 (empty cell below "12. Name")
            'role': (2, 2),              # Row 3, Col 3 (empty cell below "13. Role")  
            'total_experience': (3, 4),   # Row 4, Col 5 (empty cell below "a. Total")
            'firm_experience': (3, 6),    # Row 4, Col 7 (empty cell below "b. With current Firm")
            'firm_location': (5, 0),      # Row 6, Col 1 (empty cell below "15. Firm name")
            'education': (7, 0),          # Row 8, Col 1 (empty cell below "16. EDUCATION")
            'registration': (7, 4),       # Row 8, Col 5 (empty cell below "17. REGISTRATION")
            'qualifications': (9, 0)      # Row 10, Col 1 (empty cell below "18. OTHER QUALIFICATIONS")
        }
        
        # Project sections based on template analysis (rows 14, 19, 24, 29 have project content)
        self.project_sections = [
            {'title_row': 13, 'year_row': 14, 'desc_row': 15},  # Project 1 (around row 14)
            {'title_row': 18, 'year_row': 19, 'desc_row': 20},  # Project 2 (around row 19)  
            {'title_row': 23, 'year_row': 24, 'desc_row': 25},  # Project 3 (around row 24)
            {'title_row': 28, 'year_row': 29, 'desc_row': 30}   # Project 4 (around row 29)
        ]
    
    def find_cell_by_pattern(self, table, patterns: List[str]) -> Optional[Tuple[int, int]]:
        """
        Find a table cell that matches any of the given patterns
        
        Args:
            table: Word table object
            patterns (List[str]): Regex patterns to match
            
        Returns:
            Optional[Tuple[int, int]]: (row_index, col_index) if found, None otherwise
        """
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.lower().strip()
                for pattern in patterns:
                    if re.search(pattern, cell_text, re.IGNORECASE):
                        return (row_idx, col_idx)
        return None
    
    def find_data_cell_near_label(self, table, label_row: int, label_col: int) -> Optional[Tuple[int, int]]:
        """
        Find the data cell near a label cell - looks for empty/white cells only
        
        Args:
            table: Word table object
            label_row (int): Row index of the label
            label_col (int): Column index of the label
            
        Returns:
            Optional[Tuple[int, int]]: (row_index, col_index) of data cell
        """
        # Function to check if a cell is a data cell (empty or minimal content)
        def is_data_cell(row_idx, col_idx):
            if row_idx >= len(table.rows) or col_idx >= len(table.rows[row_idx].cells):
                return False
            cell = table.rows[row_idx].cells[col_idx]
            cell_text = cell.text.strip()
            # Consider it a data cell if it's empty or doesn't contain field numbers/labels
            has_field_number = bool(re.search(r'\d+\.', cell_text))  # Looks for patterns like "12.", "13.", etc.
            has_field_keywords = any(keyword in cell_text.lower() for keyword in ['name', 'role', 'education', 'registration', 'experience', 'firm', 'project', 'year', 'title', 'description'])
            return len(cell_text) == 0 or (not has_field_number and not has_field_keywords)
        
        # Strategy 1: Look for empty cells to the right in the same row
        for col_offset in range(1, min(4, len(table.rows[label_row].cells) - label_col)):
            target_col = label_col + col_offset
            if is_data_cell(label_row, target_col):
                print(f"    Found data cell at row {label_row}, col {target_col} (right of label)")
                return (label_row, target_col)
        
        # Strategy 2: Look for empty cells below in the same column
        for row_offset in range(1, min(3, len(table.rows) - label_row)):
            target_row = label_row + row_offset
            if is_data_cell(target_row, label_col):
                print(f"    Found data cell at row {target_row}, col {label_col} (below label)")
                return (target_row, label_col)
        
        # Strategy 3: Look for empty cells in nearby positions (below and to the right)
        for row_offset in range(1, min(3, len(table.rows) - label_row)):
            for col_offset in range(1, min(4, len(table.rows[label_row + row_offset].cells) - label_col)):
                target_row = label_row + row_offset
                target_col = label_col + col_offset
                if is_data_cell(target_row, target_col):
                    print(f"    Found data cell at row {target_row}, col {target_col} (diagonal from label)")
                    return (target_row, target_col)
        
        print(f"    No suitable data cell found near label at row {label_row}, col {label_col}")
        return None
    
    def smart_text_fill(self, cell, text: str, max_length: int = 100):
        """
        Intelligently fill a cell with text, handling line breaks and length
        
        Args:
            cell: Word table cell object
            text (str): Text to insert
            max_length (int): Maximum length before line breaking
        """
        if not text:
            return
        
        # Safety check: Don't overwrite cells that contain field labels
        existing_text = cell.text.strip()
        if existing_text and any(pattern in existing_text.lower() for pattern in ['name', 'role', 'education', 'registration', 'qualifications', 'experience', 'firm', 'project']):
            print(f"    ⚠️ Skipping cell with existing label content: {existing_text[:30]}...")
            return
        
        # Clear existing content only if it's safe to do so
        cell.text = ""
        
        # Split long text into multiple lines
        if len(text) > max_length:
            words = text.split()
            lines = []
            current_line = []
            current_length = 0
            
            for word in words:
                if current_length + len(word) + 1 <= max_length:
                    current_line.append(word)
                    current_length += len(word) + 1
                else:
                    if current_line:
                        lines.append(" ".join(current_line))
                    current_line = [word]
                    current_length = len(word)
            
            if current_line:
                lines.append(" ".join(current_line))
            
            # Add lines to cell
            for i, line in enumerate(lines):
                if i > 0:
                    self.add_line_break(cell)
                self.add_run_to_cell(cell, line)
        else:
            self.add_run_to_cell(cell, text)
    
    def add_line_break(self, cell):
        """Add a line break to a cell"""
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[-1] if paragraph.runs else paragraph.add_run()
        run.add_break()
    
    def add_run_to_cell(self, cell, text: str):
        """Add a text run to a cell with proper formatting"""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        run.font.size = Pt(9)  # Smaller font for better fitting
        return run
    
    def fill_employee_data(self, table, employee_data: Dict):
        """
        Fill employee data into the table using precise cell mapping
        
        Args:
            table: Word table object
            employee_data (Dict): Employee data dictionary
        """
        print("🔍 Filling employee information...")
        
        # Map of field names to data values
        field_mapping = {
            'name': employee_data.get('name', ''),
            'role': employee_data.get('role_in_contract', ''),
            'total_experience': str(employee_data.get('years_experience', {}).get('total', '')),
            'firm_experience': str(employee_data.get('years_experience', {}).get('with_current_firm', '')),
            'firm_location': employee_data.get('firm_name_and_location', ''),
            'education': employee_data.get('education', ''),
            'registration': employee_data.get('current_professional_registration', ''),
            'qualifications': employee_data.get('other_professional_qualifications', '')
        }
        
        # Fill each field using precise cell coordinates
        filled_count = 0
        for field_name, value in field_mapping.items():
            if not value:
                continue
                
            # Get precise cell coordinates from mapping
            cell_pos = self.employee_cell_mapping.get(field_name)
            if cell_pos:
                row_idx, col_idx = cell_pos
                try:
                    # Verify cell exists and is safe to fill
                    if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                        cell = table.rows[row_idx].cells[col_idx]
                        
                        # Safety check: make sure it's not a header cell
                        existing_text = cell.text.strip()
                        if not existing_text or len(existing_text) < 5:  # Empty or minimal content
                            self.smart_text_fill(cell, str(value), max_length=80)
                            print(f"  ✓ {field_name}: {str(value)[:50]}... → Row {row_idx+1}, Col {col_idx+1}")
                            filled_count += 1
                        else:
                            print(f"  ⚠️ Skipping {field_name} - cell has content: {existing_text[:30]}...")
                    else:
                        print(f"  ❌ Invalid cell position for {field_name}: {cell_pos}")
                except Exception as e:
                    print(f"  ❌ Error filling {field_name}: {e}")
            else:
                print(f"  ⚠️ No cell mapping found for {field_name}")
        
        print(f"📊 Filled {filled_count} employee fields")
    
    def fill_project_data(self, table, projects: List[Dict]):
        """
        Fill project data into the table
        
        Args:
            table: Word table object
            projects (List[Dict]): List of project dictionaries
        """
        if not projects:
            print("⚠️ No projects to fill")
            return
        
        print(f"🔍 Filling {len(projects)} projects...")
        
        # Find all project-related rows by looking for project patterns
        project_rows = []
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                cell_text = cell.text.lower()
                if any(pattern in cell_text for pattern in ['title.*location', 'year.*completed', 'description', 'brief.*scope']):
                    if row_idx not in project_rows:
                        project_rows.append(row_idx)
        
        print(f"📍 Found potential project rows: {project_rows}")
        
        # Fill projects into available rows
        filled_projects = 0
        for i, project in enumerate(projects[:10]):  # Limit to 10 projects
            if i >= len(project_rows) // 4:  # Assuming 4 rows per project
                break
            
            base_row = project_rows[i * 4] if i * 4 < len(project_rows) else project_rows[-1]
            
            try:
                # Use precise project section mapping
                if i < len(self.project_sections):
                    section = self.project_sections[i]
                    
                    # Project title and location
                    title = project.get('title_and_location', '')
                    if title:
                        title_row = section['title_row']
                        if title_row < len(table.rows):
                            # Find empty cell in title row (usually column 1)
                            for col_idx in range(1, min(7, len(table.rows[title_row].cells))):
                                cell = table.rows[title_row].cells[col_idx]
                                if not cell.text.strip():
                                    self.smart_text_fill(cell, title, max_length=80)
                                    print(f"  ✓ Project {i+1} title: {title[:40]}... → Row {title_row+1}, Col {col_idx+1}")
                                    break
                    
                    # Project year (both professional services and construction)
                    year_data = project.get('year_completed', {})
                    if isinstance(year_data, dict):
                        prof_year = year_data.get('professional_services', '')
                        const_year = year_data.get('construction', '')
                        year_str = prof_year or const_year or ''
                    else:
                        year_str = str(year_data) if year_data else ''
                    
                    if year_str:
                        year_row = section['year_row']
                        if year_row < len(table.rows):
                            # Professional services year (column 4-5)
                            for col_idx in range(4, min(6, len(table.rows[year_row].cells))):
                                cell = table.rows[year_row].cells[col_idx]
                                if not cell.text.strip():
                                    self.smart_text_fill(cell, year_str, max_length=20)
                                    print(f"  ✓ Project {i+1} year: {year_str} → Row {year_row+1}, Col {col_idx+1}")
                                    break
                    
                    # Combined project description (description + role + cost)
                    description_parts = []
                    
                    # Get description
                    desc_data = project.get('description', {})
                    if isinstance(desc_data, dict):
                        scope = desc_data.get('scope', '')
                        if scope:
                            description_parts.append(scope)
                    elif desc_data:
                        description_parts.append(str(desc_data))
                    
                    # Get role
                    role = desc_data.get('role', '') if isinstance(desc_data, dict) else ''
                    if role:
                        description_parts.append(f"Role: {role}")
                    
                    # Get cost
                    cost = desc_data.get('cost', '') if isinstance(desc_data, dict) else ''
                    if cost:
                        description_parts.append(f"Cost: {cost}")
                    
                    # Combine all parts
                    combined_description = ". ".join(description_parts)
                    
                    if combined_description:
                        desc_row = section['desc_row']
                        if desc_row < len(table.rows):
                            # Find empty cell for description (spans multiple columns)
                            for col_idx in range(1, min(7, len(table.rows[desc_row].cells))):
                                cell = table.rows[desc_row].cells[col_idx]
                                if not cell.text.strip():
                                    self.smart_text_fill(cell, combined_description, max_length=200)
                                    print(f"  ✓ Project {i+1} description: {combined_description[:50]}... → Row {desc_row+1}, Col {col_idx+1}")
                                    break
                    
                    # Handle checkbox for "performed with same firm"
                    performed_with_same_firm = project.get('performed_with_same_firm', False)
                    print(f"  🔍 Project {i+1} checkbox data: {performed_with_same_firm} (type: {type(performed_with_same_firm)})")
                    
                    if performed_with_same_firm is not None:
                        checkbox_symbol = "☑" if performed_with_same_firm else "☐"
                        checkbox_text = f"{checkbox_symbol} Performed with same firm"
                        
                        # Try to find a cell near the description for the checkbox
                        # Look in the row after description or in the last column
                        checkbox_row = min(desc_row + 1, len(table.rows) - 1)
                        for col_idx in range(5, min(7, len(table.rows[checkbox_row].cells))):
                            cell = table.rows[checkbox_row].cells[col_idx]
                            if not cell.text.strip():
                                self.smart_text_fill(cell, checkbox_text, max_length=30)
                                print(f"  ✓ Project {i+1} checkbox: {checkbox_symbol} → Row {checkbox_row+1}, Col {col_idx+1}")
                                break
                
                filled_projects += 1
                
            except Exception as e:
                print(f"❌ Error processing project {i+1}: {e}")
        
        print(f"📊 Filled {filled_projects} projects")
    
    def generate_section_e_docx(self, employee_data: Dict, output_filename: Optional[str] = None) -> str:
        """
        Generate Section E DOCX using template
        
        Args:
            employee_data (Dict): Employee data dictionary
            output_filename (Optional[str]): Custom output filename
            
        Returns:
            str: Path to generated DOCX file
        """
        # Get template path (check local first, then download from Supabase if needed)
        if os.path.exists(self.template_path):
            # Use local template
            actual_template_path = self.template_path
            print(f"📄 Using local template: {actual_template_path}")
        else:
            # Try to get template using downloader (for Supabase cloud templates)
            template_filename = os.path.basename(self.template_path)
            actual_template_path = self.template_downloader.get_template_path(template_filename)
            if not actual_template_path:
                raise FileNotFoundError(f"Template not found locally or in cloud: {self.template_path}")
        
        # Generate output filename
        if not output_filename:
            safe_name = "".join(c for c in employee_data.get('name', 'Unknown') if c.isalnum() or c in (' ', '-', '_')).strip()
            output_filename = f"SectionE_{safe_name.replace(' ', '_')}.docx"
        
        output_path = self.output_dir / output_filename
        
        try:
            # Load template
            print(f"📄 Loading template: {actual_template_path}")
            doc = Document(actual_template_path)
            
            # Find the main table (should be the largest one)
            if not doc.tables:
                raise Exception("No tables found in template")
            
            main_table = max(doc.tables, key=lambda t: len(t.rows) * len(t.columns))
            print(f"📋 Using table with {len(main_table.rows)} rows × {len(main_table.columns)} columns")
            
            # Fill employee data
            self.fill_employee_data(main_table, employee_data)
            
            # Fill project data
            projects = employee_data.get('relevant_projects', [])
            if projects:
                self.fill_project_data(main_table, projects)
            
            # Save the document
            doc.save(str(output_path))
            print(f"✅ DOCX generated successfully: {output_path}")
            
            return str(output_path)
            
        except Exception as e:
            print(f"❌ Error generating DOCX: {e}")
            raise
    
    def process_json_file(self, json_path: str) -> List[str]:
        """
        Process all employees from a JSON file
        
        Args:
            json_path (str): Path to JSON file with employee data
            
        Returns:
            List[str]: List of generated file paths
        """
        if not os.path.exists(json_path):
            raise FileNotFoundError(f"JSON file not found: {json_path}")
        
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        employees = data.get('resumes', [])
        if not employees:
            raise ValueError("No employee data found in JSON file")
        
        generated_files = []
        print(f"🚀 Processing {len(employees)} employees...")
        
        for i, employee in enumerate(employees, 1):
            try:
                employee_data = employee.get('data', employee)  # Handle both formats
                name = employee_data.get('name', f'Employee_{i}')
                print(f"\n📝 Processing {i}/{len(employees)}: {name}")
                
                output_path = self.generate_section_e_docx(employee_data)
                generated_files.append(output_path)
                
            except Exception as e:
                print(f"❌ Failed to process employee {i}: {e}")
        
        return generated_files


def main():
    """Main function"""
    import argparse
    
    parser = argparse.ArgumentParser(
        description="Generate Section E resumes using DOCX templates"
    )
    parser.add_argument("--template", default="templates/Template Section E Resume.docx",
                       help="Path to DOCX template file")
    parser.add_argument("--data", help="Path to JSON file with employee data")
    parser.add_argument("--employee", help="Generate for specific employee name")
    parser.add_argument("--output", help="Output filename (for single employee)")
    
    args = parser.parse_args()
    
    try:
        generator = DOCXSectionEGenerator(args.template)
        
        if args.data:
            # Process JSON file
            generated_files = generator.process_json_file(args.data)
            print(f"\n🎉 Successfully generated {len(generated_files)} DOCX files!")
            print("📁 Output directory:", generator.output_dir)
            
        elif args.employee:
            # Generate for specific employee (would need integration with database)
            print("❌ Employee-specific generation not yet implemented")
            print("💡 Use --data flag with JSON file instead")
            
        else:
            print("❌ Please specify --data (JSON file) or --employee (name)")
            print("💡 Example: python3 docx_section_e_generator.py --data ParsedFiles/results.json")
    
    except Exception as e:
        print(f"❌ Error: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main() 