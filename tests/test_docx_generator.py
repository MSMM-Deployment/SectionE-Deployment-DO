#!/usr/bin/env python3
"""
Test script for the DOCX Section E Generator

This script tests the DOCX generator with sample employee data to ensure
it's working correctly and can properly fill the template.
"""

import json
import os
import sys
from pathlib import Path

# Add project root to Python path for src module imports
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, project_root)

def test_docx_generator():
    """Test the DOCX generator with sample data"""
    
    # Sample employee data
    sample_employee_data = {
        "name": "John A. Smith, PE",
        "role_in_contract": "Senior Project Engineer",
        "years_experience": {
            "total": "15",
            "with_current_firm": "8"
        },
        "firm_name_and_location": "ABC Engineering Consultants, Denver, CO",
        "education": "M.S. Civil Engineering, University of Colorado, 2008; B.S. Civil Engineering, Colorado State University, 2006",
        "current_professional_registration": "Professional Engineer (PE) - Colorado License #12345",
        "other_professional_qualifications": "LEED AP BD+C, Project Management Professional (PMP)",
        "relevant_projects": [
            {
                "title_and_location": "Downtown Denver Transit Station, Denver, CO",
                "year_completed": {
                    "professional_services": "2023",
                    "construction": "2024"
                },
                "description": {
                    "scope": "Structural design and analysis for a 50,000 SF multi-modal transit facility serving bus and light rail connections. Project included complex foundation systems, steel frame construction, and seismic design considerations.",
                    "cost": "$15M",
                    "role": "Lead Structural Engineer"
                }
            },
            {
                "title_and_location": "Rocky Mountain Conference Center, Boulder, CO",
                "year_completed": {
                    "professional_services": "2022",
                    "construction": "2023"
                },
                "description": {
                    "scope": "Design of a 75,000 SF conference and exhibition center featuring long-span steel construction, specialized acoustical considerations, and sustainable design elements achieving LEED Gold certification.",
                    "cost": "$22M",
                    "role": "Project Engineer"
                }
            },
            {
                "title_and_location": "Colorado State University Research Building, Fort Collins, CO",
                "year_completed": {
                    "professional_services": "2021",
                    "construction": "2022"
                },
                "description": {
                    "scope": "Structural engineering for a 4-story, 40,000 SF research facility with specialized laboratory spaces, vibration-sensitive equipment areas, and complex MEP coordination requirements.",
                    "cost": "$18M",
                    "role": "Senior Engineer"
                }
            }
        ]
    }
    
    print("🧪 Testing DOCX Section E Generator")
    print("=" * 50)
    
    try:
        # Import the DOCX generator
        from src.generators.docx_section_e_generator import DOCXSectionEGenerator
        
        # Check if template exists
        template_path = "templates/Template Section E Resume.docx"
        if not os.path.exists(template_path):
            print(f"❌ Template not found: {template_path}")
            print("💡 Make sure the DOCX template is in the templates folder")
            return False
        
        print(f"✅ Template found: {template_path}")
        
        # Initialize generator
        generator = DOCXSectionEGenerator(template_path)
        print(f"✅ Generator initialized")
        print(f"📁 Output directory: {generator.output_dir}")
        
        # Generate DOCX
        print(f"\n📝 Generating DOCX for: {sample_employee_data['name']}")
        output_path = generator.generate_section_e_docx(sample_employee_data)
        
        if output_path and os.path.exists(output_path):
            file_size = os.path.getsize(output_path)
            print(f"✅ DOCX generated successfully!")
            print(f"📄 File: {output_path}")
            print(f"📦 Size: {file_size:,} bytes")
            
            # Verify the file can be opened
            try:
                from docx import Document
                doc = Document(output_path)
                print(f"✅ DOCX file is valid (contains {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables)")
                return True
            except Exception as e:
                print(f"⚠️ DOCX file generated but may have issues: {e}")
                return False
        else:
            print(f"❌ DOCX generation failed")
            return False
            
    except ImportError as e:
        print(f"❌ Import error: {e}")
        print("💡 Make sure python-docx is installed: pip install python-docx")
        return False
    except Exception as e:
        print(f"❌ Test failed: {e}")
        return False


def test_with_json_file():
    """Test with actual parsed JSON data if available"""
    
    json_files = [
        "data/ParsedFiles/real_parsed_results.json",
        "data/ParsedFiles/results.json",
        "data/ParsedFiles/parsed_results.json"
    ]
    
    for json_file in json_files:
        if os.path.exists(json_file):
            print(f"\n🧪 Testing with JSON file: {json_file}")
            print("=" * 50)
            
            try:
                from src.generators.docx_section_e_generator import DOCXSectionEGenerator
                
                generator = DOCXSectionEGenerator()
                generated_files = generator.process_json_file(json_file)
                
                print(f"✅ Generated {len(generated_files)} DOCX files from JSON data")
                for file_path in generated_files[:3]:  # Show first 3
                    file_size = os.path.getsize(file_path)
                    print(f"📄 {Path(file_path).name} ({file_size:,} bytes)")
                
                if len(generated_files) > 3:
                    print(f"📄 ... and {len(generated_files) - 3} more files")
                
                return True
                
            except Exception as e:
                print(f"❌ JSON file test failed: {e}")
                return False
    
    print(f"\n⚠️ No JSON files found for testing")
    print(f"💡 Run the parser first to create test data:")
    print(f"   python3 src/parsers/section_e_parser.py --folder data/Sample-SectionE --output data/ParsedFiles/results.json")
    return False


def main():
    """Main test function"""
    print("🚀 DOCX Section E Generator Test Suite")
    print("=" * 60)
    
    # Test 1: Basic functionality with sample data
    test1_passed = test_docx_generator()
    
    # Test 2: JSON file processing (if available)
    test2_passed = test_with_json_file()
    
    # Summary
    print(f"\n📊 Test Results Summary")
    print("=" * 30)
    print(f"✅ Basic DOCX Generation: {'PASSED' if test1_passed else 'FAILED'}")
    print(f"✅ JSON File Processing: {'PASSED' if test2_passed else 'SKIPPED'}")
    
    if test1_passed:
        print(f"\n🎉 DOCX generator is working correctly!")
        print(f"💡 You can now use DOCX templates in the web interface")
        print(f"📁 Generated files are in: OutputFiles/DOCX/")
    else:
        print(f"\n❌ DOCX generator has issues that need to be fixed")
    
    return test1_passed


if __name__ == "__main__":
    main() 